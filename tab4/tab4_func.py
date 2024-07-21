import gradio as gr
import os
from docx import Document
import codecs
import re


#vttファイルのタイムスタンプ桁数を統一。


def unify_timestamps_vtt(text):
    # Define patterns for different timestamp formats
    pattern_1_digit = re.compile(r'(\d{2}:\d{2}:\d{2}\.\d)(?!\d)')
    pattern_2_digits = re.compile(r'(\d{2}:\d{2}:\d{2}\.\d{2})(?!\d)')
    pattern_3_digit = re.compile(r'(\d{1}:\d{2}:\d{2}\.\d)(?!\d)')
    pattern_4_digits = re.compile(r'(\d{1}:\d{2}:\d{2}\.\d{2})(?!\d)')

    # Replace 1-digit and 2-digit millisecond formats with 3-digit format for HH:MM:SS format
    if pattern_1_digit.search(text):
        text = pattern_1_digit.sub(lambda x: x.group(1) + '00', text)
        print("pattern1")
    if pattern_2_digits.search(text):
        text = pattern_2_digits.sub(lambda x: x.group(1) + '0', text)
        print("pattern2")

    # Replace 1-digit and 2-digit millisecond formats with 3-digit format for H:MM:SS format
    if pattern_3_digit.search(text):
        text = pattern_3_digit.sub(lambda x: x.group(1) + '00', text)
        print("pattern3")
    if pattern_4_digits.search(text):
        text = pattern_4_digits.sub(lambda x: x.group(1) + '0', text)
        print("pattern4")

    return text


#srtファイルのタイムスタンプ桁数を統一。
def unify_timestamps(text):
    # Define patterns for different timestamp formats
    pattern_1_digit = re.compile(r'(\d{2}:\d{2}:\d{2}\,\d)(?!\d)')
    pattern_2_digits = re.compile(r'(\d{2}:\d{2}:\d{2}\,\d{2})(?!\d)')
    
    # Replace 1-digit millisecond format with 3-digit format
    text = pattern_1_digit.sub(lambda x: x.group(1) + '00', text)
    
    # Replace 2-digit millisecond format with 3-digit format
    text = pattern_2_digits.sub(lambda x: x.group(1) + '0', text)
    
    return text
def unify_timestamps_forlist(lines, format_type):
    if format_type == 'vtt':
        unify_timestamps_func = unify_timestamps_vtt
    elif format_type == 'srt':
        unify_timestamps_func = unify_timestamps
    else:
        raise ValueError(f"Unsupported format type: {format_type}")

    return [unify_timestamps_func(line) if '-->' in line else line for line in lines]

def read_file_content(file):
    if file is None:
        return """<div style='color: orange !important; font-family: inherit; text-align: center; 
                display: flex; align-items: flex-start; justify-content: center; height: 400px; padding-top: 40px;'>
                No file uploaded
                </div>"""

    file_extension = os.path.splitext(file)[1]
    content = ""
    # file.nameをfileに変更　↑

    '''    if file_extension == '.docx':
            doc = Document(file)#file.nameのname除去中
            content = "\n".join([para.text for para in doc.paragraphs])
            content = f"""<pre style="white-space: pre-wrap; overflow-y: auto; height: 400px; word-wrap: break-word; padding: 10px; font-family: inherit; font-size: inherit;">{content}</pre>"""
    '''
    if file_extension == '.txt':#file.nameの.name除去中
        with codecs.open(file, 'r', 'utf-8') as f:
            content = f.read()
            content = f"""<pre style="white-space: pre-wrap; overflow-y: auto; height: 500px; word-wrap: break-word; padding: 10px; font-family: inherit; font-size: inherit;">{content}</pre>"""

    elif file_extension == '.srt' : # file.nameのname除去中。
        with codecs.open(file, 'r', 'utf-8') as f:
            content = f.read()
            content=unify_timestamps(content)
            content = f"""<pre style="white-space: pre-wrap; overflow-y: auto; height: 500px; word-wrap: break-word; padding: 10px; font-family: inherit; font-size: inherit;">{content}</pre>"""

    elif file_extension == '.vtt': # file.nameのname除去中。
        with codecs.open(file,'r',encoding='utf-8')as f:
            content = f.read()
            content = unify_timestamps_vtt(content)
            content = f"""<pre style="white-space: pre-wrap; overflow-y: auto; height: 500px; word-wrap: break-word; padding: 10px; font-family: inherit; font-size: inherit;">{content}</pre>"""
    else:
        return None
    return content

def display_file_content(file):
    if file is None:
        return read_file_content(file), gr.update(value=""), gr.update(value=None)
    return read_file_content(file), gr.update(), gr.update()

def save_translated_content(file, translated_text):
    if file==None:
        return []
    
    file_name, file_extension = os.path.splitext(file) #file.nameの.name除去中
    output_file_path = file_name + "_ja" + file_extension

    if file_extension == '.docx':
        doc = Document()
        doc.add_paragraph(translated_text)
        doc.save(output_file_path)

    elif file_extension == '.txt':
        with codecs.open(output_file_path, 'w', 'utf-8') as f:
            f.write(translated_text)
    
    elif  file_extension == '.srt' or file_extension == '.vtt':
        content = re.sub(r'[\u200B-\u200D\uFEFF]', '', translated_text)
        content = re.sub(r'\s+', '', content)
        content = content.replace('\u200B', '')

        if file_extension == '.srt':
            pattern = re.compile(r'(\d{1,4})(\d{2}:\d{2}:\d{2},\d{3}-->\d{2}:\d{2}:\d{2},\d{3})')
        elif file_extension == '.vtt':
            pattern = re.compile(r'(\d{1,4})(\d{1}:\d{2}:\d{2}.\d{3}-->\d{1}:\d{2}:\d{2}.\d{3})')
        matches = pattern.findall(content)      
        print(f"d1(len(matches)):{len(matches)}")  
        if len(matches)==0:
            pattern = re.compile(r'(\d{1,4})(\d{2}:\d{2}:\d{2}.\d{3}-->\d{2}:\d{2}:\d{2}.\d{3})')
            matches = pattern.findall(content)
            print(f"d2(len(matches)):{len(matches)}")  
        segments = pattern.split(content)
        corrected_content = []
        
        for i in range(1, len(segments), 3):
            segment_id = segments[i]
            timestamp = segments[i + 1]
            text = segments[i + 2]
            
            corrected_content.append(f"{segment_id}")
            corrected_content.append(timestamp.replace('-->', ' --> '))
            corrected_content.append(text)

        # Ensure each subtitle block is separated by exactly one empty line
        if file_extension == '.vtt':
            corrected_content ="WEBVTT\n\n"+"\n\n".join("\n".join(block) for block in zip(*[iter(corrected_content)]*3))        
        else:
            corrected_content = "\n\n".join("\n".join(block) for block in zip(*[iter(corrected_content)]*3))        
        with open(output_file_path,'w',encoding='utf-8') as file:
            file.write(corrected_content+'\n')



    return output_file_path

def translate(file, translated_text):
    return save_translated_content(file, translated_text)
