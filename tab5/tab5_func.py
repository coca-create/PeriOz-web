import gradio as gr
import docx
from docx import Document
import tempfile
import os
import re
import zipfile
from tab4 import tab4_func as t4


def webvtt_remover_NR(sentence):
    sentence = re.sub(r'[\u200B-\u200D\uFEFF]', '', sentence)
    sentence = sentence.replace('\r\n', '\n').replace('\r', '\n')
    sentence = sentence.replace('\u200B',"")

    pattern1 = re.compile(r'(\d+)(\d{1}:\d{2}:\d{2}\.\d{3}-->\d{1}:\d{2}:\d{2}\.\d{3})', re.DOTALL)
    pattern2 = re.compile(r'(\d+)(\d{2}:\d{2}:\d{2}\.\d{3}-->\d{2}:\d{2}:\d{2}\.\d{3})', re.DOTALL)

    match = pattern1.search(sentence)
    print(f"webvtt_remover_match[search]_pattern1(d1):{match}")
    if match:
        start_index = match.start(1)
        rm_webvtt_sentence = sentence[start_index:]
    else:
        match = pattern2.search(sentence)
        print(f"webvtt_remover_match[search]_pattern2(d2):{match}")
        if match:
            start_index = match.start(1)
            rm_webvtt_sentence = sentence[start_index:]
        else:
            rm_webvtt_sentence = sentence

    return rm_webvtt_sentence
'''def webvtt_remover_NR(sentence):
    sentence = re.sub(r'[\u200B-\u200D\uFEFF]', '', sentence)
    sentence = sentence.replace('\r\n', '\n').replace('\r', '\n')
    sentence = sentence.replace('\u200B',"")

    pattern = re.compile(r'(\d+)(\d{1}:\d{2}:\d{2}\.\d{3}-->\d{1}:\d{2}:\d{2}\.\d{3})', re.DOTALL)
    match = pattern.search(sentence)
    if match:
        start_index = match.start(1)
        rm_webvtt_sentence = sentence[start_index:]
    elif pattern.search(re.compile(r'(\d+)(\d{2}:\d{2}:\d{2}\.\d{3} --> \d{2}:\d{2}:\d{2}\.\d{3})', re.DOTALL)
        ):
        match =pattern.search(re.compile(r'(\d+)(\d{2}:\d{2}:\d{2}\.\d{3} --> \d{2}:\d{2}:\d{2}\.\d{3})', re.DOTALL)
        )
        start_index = match.start(1)
        rm_webvtt_sentence = sentence[start_index]  
    else:
        rm_webvtt_sentence = sentence
    return rm_webvtt_sentence'''

def convert_docx_to_srttxt(docx_files):
    output_files = []
    if docx_files is None:
        return []
    for docx_file in docx_files:
        try:
            filename = os.path.basename(docx_file)
            base_name, ext = os.path.splitext(filename)
            clean_name = re.sub(r'[\r\n]+', '', base_name)
            print(f"Processing file: {clean_name}")

            if clean_name.endswith('_srt 1'):
                output_filename = clean_name.replace('_srt 1', '_ja.srt')
            elif clean_name.endswith("_srt"):
                output_filename = clean_name.replace("_srt", "_ja.srt")
            elif clean_name.endswith('_vtt 1'):
                output_filename = clean_name.replace('_vtt 1', '_ja.vtt')
            elif clean_name.endswith("_vtt"):
                output_filename = clean_name.replace("_vtt", "_ja.vtt")
            elif clean_name.endswith('_txtnr 1'):
                output_filename = clean_name.replace('_txtnr 1', '_NR_ja.txt')
            elif clean_name.endswith("_txtnr"):
                output_filename = clean_name.replace("_txtnr", "_NR_ja.txt")
            elif clean_name.endswith('_txtr 1'):
                output_filename = clean_name.replace('_txtr 1', '_R_ja.txt')
            elif clean_name.endswith("_txtr"):
                output_filename = clean_name.replace("_txtr", "_R_ja.txt")
            else:
                print(f"Skipping file with unrecognized pattern: {clean_name}")
                continue

            doc = Document(docx_file)
            content = "\n".join([para.text for para in doc.paragraphs])
            print(f"Initial content read from file: {content[:100]}")  # Show only the first 100 characters for brevity

            if output_filename.endswith('.srt'):
                content = t4.unify_timestamps(content)
                content = re.sub(r'\s+', '', content)
                content = content.replace('\u200B', '')
                pattern = re.compile(r'(\d{1,4})(\d{2}:\d{2}:\d{2},\d{3}-->\d{2}:\d{2}:\d{2},\d{3})')

            elif output_filename.endswith('.vtt'):
                content = t4.unify_timestamps_vtt(content)
                
                content = re.sub(r'\s+', '', content)
                content = content.replace('\u200B', '')
                content = webvtt_remover_NR(content)
                print(content)
                pattern_1 = re.compile(r'(\d{1,4})(\d{1}:\d{2}:\d{2}\.\d{3}-->\d{1}:\d{2}:\d{2}\.\d{3})')
                pattern_2 = re.compile(r'(\d{1,4})(\d{2}:\d{2}:\d{2}\.\d{3}-->\d{2}:\d{2}:\d{2}\.\d{3})')
                match = pattern_1.search(content)
                if match:
                    pattern = pattern_1
                    print(f"pattern1_match:{match}")
                else:
                    pattern = pattern_2
                    print(f"pattern1 doesn't match. so... pattern2")

            if output_filename.endswith('.srt') or output_filename.endswith('.vtt'):
                segments = pattern.split(content)
                corrected_content = []
                for i in range(1, len(segments), 3):
                    segment_id = segments[i]
                    timestamp = segments[i + 1]
                    text = segments[i + 2]
                    
                    corrected_content.append(f"{segment_id}")
                    corrected_content.append(timestamp.replace('-->', ' --> '))
                    corrected_content.append(text)


            if output_filename.endswith('.vtt'):
                final_content = 'WEBVTT\n\n' + "\n\n".join("\n".join(block) for block in zip(*[iter(corrected_content)] * 3))
            if output_filename.endswith('.srt'):
                final_content = "\n\n".join("\n".join(block) for block in zip(*[iter(corrected_content)] * 3))

            output_filepath = os.path.join(tempfile.gettempdir(), output_filename)
            if output_filename.endswith('.txt'):
                final_content = content

            with open(output_filepath, 'w', encoding='utf-8') as output_file:
                output_file.write(final_content)
            
            output_files.append(output_filepath)
        except Exception as e:
            print(f"An error occurred while processing {filename}: {str(e)}")

    if len(output_files) > 1:
        zip_filename = os.path.join(tempfile.gettempdir(), "converted_from_docx_ja.zip")
        with zipfile.ZipFile(zip_filename, 'w') as zipf:
            for file in output_files:
                zipf.write(file, os.path.basename(file))
        
        output_files.append(zip_filename)
    
    return output_files

def clear_inputs():
    return None, None



def process_doc_files(files):
    output_files = []
    if files is None:
        return []
    for file in files:
        filename = os.path.basename(file)
        print(filename)
        match = re.match(r"(.+?)(_NR\.txt|_R\.txt|\.srt|\.vtt)$", filename)
        if not match:
            continue  # skip files with unknown extensions
        
        basename, ext = match.groups()
        if ext == ".srt":
            doc_filename = f"{basename}_srt.docx"
        elif ext == ".vtt":
            doc_filename = f"{basename}_vtt.docx"
        elif ext == "_NR.txt":
            doc_filename = f"{basename}_txtnr.docx"
        elif ext == "_R.txt":
            doc_filename = f"{basename}_txtr.docx"

        if ext in ['.srt', '.vtt']:
            with open(file, 'r', encoding='utf-8') as f:
                content = f.read()
            print(content)
            
            if ext == '.srt':
                unified_content = t4.unify_timestamps(content)
            else:
                unified_content = t4.unify_timestamps_vtt(content)

            doc = Document()
            doc.add_paragraph(unified_content)
            doc.save(doc_filename)
            output_files.append(doc_filename)

        elif ext in ["_NR.txt", "_R.txt"]:
            with open(file.name, 'r', encoding='utf-8') as f:
                content = f.read()
            doc = Document()
            doc.add_paragraph(content)
            doc.save(doc_filename)
            output_files.append(doc_filename)
    
    if len(output_files) > 1:
        zip_filename = "converted_from_srttxt_en.zip"
        with zipfile.ZipFile(zip_filename, 'w') as zipf:
            for file in output_files:
                zipf.write(file)
        output_files.append(zip_filename)
    
    return output_files

def clear_both():
    return None, None


